from django.shortcuts import render,redirect
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt,csrf_protect
import json
import os
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
import google.generativeai as genai
import os
from dotenv import load_dotenv
import json
import random
##User authentication imports
from django.contrib.auth.models import User
from django.contrib.auth import authenticate, login,logout
from django.contrib.auth.decorators import login_required
################################################################
##rag-libraries##
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from langchain_pinecone import PineconeEmbeddings
from pinecone import Pinecone, ServerlessSpec
import time
from langchain_pinecone import PineconeVectorStore
from langchain_openai import ChatOpenAI
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain import hub
from langchain_openai.embeddings import OpenAIEmbeddings
from langchain_core.prompts import ChatPromptTemplate
###########################################################
#doc generation
from django.shortcuts import render, redirect
from django.http import HttpResponse
from .forms import WillForm, LicenseForm, LoanAgreementForm, DeedOfHypothecationForm, BailBondForm, ContractBondForm, SimpleMoneyBondForm, EmployeeBondForm, MortgageDeedForm, RentAgreementForm, SaleAgreementForm, BailPetitionForm, DeedOfAdoptionForm, LeaveAndLicenseAgreementForm
import docx
from django.shortcuts import render
################################################################
genai.configure(api_key=os.environ.get('GOOGLE_API_KEY'))
OPENAI_API_KEY=os.environ.get("OPENAI_API_KEY") 
model = genai.GenerativeModel("gemini-1.5-flash")
openaiEmbedding=OpenAIEmbeddings(model="text-embedding-ada-002")
pc=Pinecone(api_key=os.environ.get('PINECONE_API_KEY'))
spec=ServerlessSpec(cloud='aws',region='us-east-1')
index_name = "rag-bot-index"+str(random.randint(1,1000))
namespace = "rag-rohith"
prompt = ChatPromptTemplate.from_messages([
    ("system", "You are a specialized Indian Legal Expert chatbot designed to provide concise, context-specific responses related to Indian laws, including the Indian Penal Code and other relevant legislations. Use only the provided context to answer questions,context='{context}', and if the context is not sufficient, respond with: 'Sorry, I couldn't find an answer related to the uploaded document.'"),
    ("human", "{input}")
])

file_name=""
retrieval_chain=None
combine_docs_chain=None
# Create your views here.
load_dotenv()
def home(request):
    return render(request, 'core/home.html')

def features(request):
    return render(request, 'core/features.html')
@login_required
def ragbot(request):
    return render(request, 'core/ragbot.html')
@login_required
def legalbot(request):
    return render(request,"core/legalbot.html")
@csrf_exempt
def login_page(request):
    return render(request, 'core/login.html')
def register_page(request): 
    return render(request, 'core/signup.html')
@csrf_exempt
def register(request):
 if request.method=="POST":
            data=json.loads(request.body)
            print(data)
            user_name=data.get("username")
            #if User.objects.get(username=user_name):
             #   return JsonResponse({"message":"Registration Failed, Username Already Exists"})
            email=data.get("email")
            password=data.get("password")
            re_password=data.get("rePassword")
            print(f"username={user_name} email={email} password={password}  re_password={re_password}")
            if password==re_password:
                #store the user data in the database
                user=User.objects.create_user(username=user_name,email=email,password=password)              
                user.save()
                print(f"user name={user.username}\temailid={user.email}\tpassword={user.password}")
                #here you can use Django's built-in User model or any other database to store user data
                return JsonResponse({"message":"Registration successful",'username':user.username})
            else:
                return JsonResponse({"message":"registration failed,passwords do not match"})
@csrf_exempt
def signin(request):
    if request.method=="POST":
        data=json.loads(request.body)
        username=data.get("username")
        password=data.get("password")
        print(f'login page-->email={username}\t\tpassword={password}')
        user=authenticate(username=username,password=password)
        print(f'authenticate function-->user={user}')
        if user is not None:
            login(request,user)
            #user_name=request.user.get_username()
            return JsonResponse({"message":"Login successful"})

        else:
            return JsonResponse({"message":"Login failed, invalid credentials"})
def signout(request):
    if request.method=="POST":
            logout(request)
            return redirect("home")    
@csrf_exempt
def ragchat(request):
    global retrieval_chain
    if request.method == 'POST':
        #data=json.loads(request.body.POST.get("message"))
        #print(data)
        user_message=request.POST.get('message')
        print(user_message)
        #response=model.generate_content(f'Generate A  concise reply for the following message from user:{user_message}.STRICTLY AVOID USING SYMBOLS IN RESPONSE<MAKE SURE THE RESPONS IS CLEAN RAW FORMATTED TEXT')
        response=retrieval_chain.invoke({'input':user_message})
        answer=response['answer']
        print(answer)
        return JsonResponse({"botReply":answer})
    
@csrf_exempt
def fileHandler(request):
    global file_name
    if request.method=='POST' and request.FILES.get('file'):
        uploaded_file=request.FILES["file"]
        file_name = default_storage.save(uploaded_file.name, ContentFile(uploaded_file.read()))
        print(file_name)
        return JsonResponse({"message":"file uploaded success!"})
    

def extract_text(file_name):
    #load the pdf
    loader=PyPDFLoader(f'./media/{file_name}')
    doc=loader.load()
    ##extract the text from the pdf
    text_splitter=CharacterTextSplitter(separator='\n',chunk_size=512,chunk_overlap=150)
    chunks=text_splitter.split_documents(doc)
    print(f'number of chunks is:{len(chunks)}')
    return chunks
def create_index_pc():
    global index_name,pc,spec,namespace,openaiEmbedding
    if index_name not in pc.list_indexes().names():
        pc.create_index(name=index_name,spec=spec,metric='euclidean',dimension=1536)
        while not pc.describe_index(index_name).status['ready']:
            time.sleep(2)
    
    
def embedding_creator(chunks):
    embeddings=PineconeVectorStore.from_documents(documents=chunks,embedding=openaiEmbedding,index_name=index_name,namespace=namespace)
    time.sleep(4)
    return embeddings
def create_rag_chain(embeddings):
    llm=ChatOpenAI(api_key=OPENAI_API_KEY,model_name='gpt-4o-mini',temperature=0.3)
    retreiver=embeddings.as_retriever(search_kwargs={"k": 8})
    combine_docs_chain=create_stuff_documents_chain(llm=llm,prompt=prompt)
    retreival_chain=create_retrieval_chain(retreiver,combine_docs_chain)
    return retreival_chain
def rag_chat(retreival_chain,question):
    response=retreival_chain.invoke({'input':{question}})
    answer=(response['answer'])
    print(answer)
    return answer

def embedder(request):
    global retrieval_chain
    create_index_pc()
    chunks=extract_text(file_name)
    embeddings=embedding_creator(chunks)
    retrieval_chain=create_rag_chain(embeddings)
    return JsonResponse({"message:":"embedding successfull,retrieval chain returned"})


def select_document(request):
    return render(request, 'core/select_document.html')

def generate_document(request, doc_type):
    if doc_type not in TEMPLATES:
        return redirect('select_document')
    
    FormClass = FORMS[doc_type]
    
    if request.method == 'POST':
        form = FormClass(request.POST)
        if form.is_valid():
            data = form.cleaned_data
            template_path = TEMPLATES.get(doc_type)
            doc = docx.Document(template_path)
            for para in doc.paragraphs:
                for key, value in data.items():
                    if value:
                        para.text = para.text.replace(f'{{{{{key}}}}}', str(value))
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename=generated_{doc_type}.docx'
            doc.save(response)
            return response
    else:
        form = FormClass()
    return render(request, 'core/generate_document.html', {'form': form, 'doc_type': doc_type})


# Define document categories and their corresponding documents
TEMPLATES = {
    'will': 'Simple-will-LawRato3.docx',
    'license': 'Licence-to-use-Copyright-LawRato2.docx',
    'loan_agreement': 'Loan-Agreement-LawRato3.docx',
    'deed_of_hypothecation': 'Deed-of-Hypothecation-HP-LawRato4.docx',
    'bail_bond': 'Bond-and-Bail-bond-under-CrPC-1973-after-Arrest-under-a-Warrant-LawRato.docx',
    'contract_bond': 'Bond-to-Secure-the-Performance-of-a-Contract-LawRato2.docx',
    'simple_money_bond': 'Simple-Money-Bond-LawRato2.docx',
    'employee_bond_for_non_compete': 'Employee-Bond-for-Non-Compete-LawRato3.docx',
    'simple_mortgage_deed': 'Simple-Mortgage-Deed-LawRato2.docx',
    'rent_agreement': 'Lease-Deed-(for-a-term-of-years)-Rent-Agreement-LawRato3.docx',
    'sale_agreement': 'Agreement-for-Sale-LawRato4.docx',
    'bail_petition' : 'Anticipatory-Bail-Petition-Format-LawRato.docx',
    'deed_of_adoption': 'Deed-of-Adoption-LawRato2.docx',
    'leave_and_license_agreement': 'Leave-and-License-Agreement-LawRato2.docx',

}

FORMS = {
    'will': WillForm,
    'license': LicenseForm,
    'loan_agreement': LoanAgreementForm,
    'deed_of_hypothecation' : DeedOfHypothecationForm,
    'bail_bond': BailBondForm,
    'contract_bond': ContractBondForm,
    'simple_money_bond': SimpleMoneyBondForm,
    'employee_bond_for_non_compete': EmployeeBondForm,
    'simple_mortgage_deed' : MortgageDeedForm,
    'rent_agreement' : RentAgreementForm,
    'sale_agreement' : SaleAgreementForm,
    'bail_petition' : BailPetitionForm,
    'deed_of_adoption' : DeedOfAdoptionForm,
    'leave_and_license_agreement': LeaveAndLicenseAgreementForm,

}

def select_document(request):
    return render(request, 'core/select_document.html')

def generate_document(request, doc_type):
    if doc_type not in TEMPLATES:
        return redirect('select_document')
    
    FormClass = FORMS[doc_type]
    
    if request.method == 'POST':
        form = FormClass(request.POST)
        if form.is_valid():
            data = form.cleaned_data
            template_path = TEMPLATES.get(doc_type)
            doc = docx.Document(template_path)
            for para in doc.paragraphs:
                for key, value in data.items():
                    if value:
                        para.text = para.text.replace(f'{{{{{key}}}}}', str(value))
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename=generated_{doc_type}.docx'
            doc.save(response)
            return response
    else:
        form = FormClass()
    return render(request, 'core/generate_document.html', {'form': form, 'doc_type': doc_type})


# Define document categories and their corresponding documents
CATEGORY_DOCUMENTS = {
    'will': {
        'Simple Will': '/will/',
    },
    'banking': {
        'Loan Agreement': '/loan_agreement/',
        'Deed Of Hypothecation' : '/deed_of_hypothecation/',

    },
    'bonds': {
        'Bail Bond': '/bail_bond/',
        'Bond to Secure Performance of a Contract': '/contract_bond/',
        'Employee Bond for Non-Compete': '/employee_bond_for_non_compete/',
        'Simple Money Bond': '/simple_money_bond/',
    },
    'contracts': {
        'Rent Agreement (for a term of years)': '/rent_agreement/',
        'Simple Mortgage Deed': '/simple_mortgage_deed/',
        'Leave and License Agreement': '/leave_and_license_agreement/',

    },
    'corporate': {
        'Agreement for Sale': '/sale_agreement/',
    },
    'criminal': {
        'Anticipatory Bail Petition Form' : '/bail_petition/',

    },
    'divorceandfamilylaw' : {
        'Deed Of Adoption' : '/deed_of_adoption/',

    },
    'trademarkandcopyright' : {
        'License to use Copyright' : '/license/',

    },

}

def category_documents(request, category_name):
    """View to display documents for a selected category."""
    documents = CATEGORY_DOCUMENTS.get(category_name, {})
    
    if not documents:
        return render(request, 'core/404.html', {"message": "Category not found"})
    
    return render(request, 'core/category_documents.html', {
        'category': category_name.replace('_', ' ').title(),
        'documents': documents
    })
#legal chatbot
from pinecone import Pinecone
from pinecone_plugins.assistant.models.chat import Message
assistant = pc.assistant.Assistant(assistant_name="legalchatbot")


def extract_content(api_response):
  """Extracts content, handling different response structures and None responses."""
  if api_response is None:
      print("Error: API Response is None.")
      return None

  if not isinstance(api_response, dict):
        print("Error: API Response is not a dictionary.")
        print("Type of API Response:", type(api_response))
        return None

  try:
    if 'message' in api_response:
      if not isinstance(api_response['message'], dict):
          print("Error: api_response['message'] is not a dictionary.")
          print("Type of api_response['message']:", type(api_response['message']))
          return None

      if 'content' in api_response['message']:
        if not isinstance(api_response['message']['content'], str):
            print("Error: api_response['message']['content'] is not a string.")
            print("Type of api_response['message']['content']:", type(api_response['message']['content']))
            return None
        content = api_response['message']['content']
        return content
      else:
        print("Error: 'content' key not found in api_response['message'].")
        return None
    else:
      print("Error: 'message' key not found in api_response.")
      return None
  except TypeError as e:
      print(f"TypeError: {e}")
      return None
  except Exception as e:
      print(f"An unexpected error occurred: {e}")
      return None

@csrf_exempt
def legalchat(request):
    if request.method == 'POST':
        #data=json.loads(request.body.POST.get("message"))
        #print(data)
        user_message=request.POST.get('usermessage-legalbot')
        print(user_message)
        msg = Message(content=f"{user_message}")
        #response=model.generate_content(f'Generate A  concise reply for the following message from user:{user_message}.STRICTLY AVOID USING SYMBOLS IN RESPONSE<MAKE SURE THE RESPONS IS CLEAN RAW FORMATTED TEXT')
        resp = dict(assistant.chat(messages=[msg]))
        print(type(resp))
        content=extract_content(resp)
        print(content)
        #print("RESPONSE IS:"+resp)
        #answer=(resp["message"]["content"])
        return JsonResponse({"botReply":content})